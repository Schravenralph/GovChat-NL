"""
Document Processing Service for Policy Scanner.

This module provides text extraction from various document formats
(PDF, HTML, DOCX) and utilities for chunking, hashing, and summarization.
"""

import hashlib
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import os

from bs4 import BeautifulSoup
from docx import Document
import pdfplumber

from open_webui.env import SRC_LOG_LEVELS

log = logging.getLogger(__name__)
log.setLevel(SRC_LOG_LEVELS.get("MAIN", "INFO"))


class ProcessingError(Exception):
    """Exception raised when document processing fails."""
    pass


class DocumentProcessor:
    """
    Service for processing policy documents.

    Supports:
    - Text extraction from PDF, HTML, DOCX
    - Content hashing for deduplication
    - Text chunking for large documents
    - Summary generation
    - Metadata enrichment
    """

    def __init__(
        self,
        max_chunk_size: Optional[int] = None,
        overlap_size: int = 200
    ):
        """
        Initialize document processor.

        Args:
            max_chunk_size: Maximum characters per chunk (defaults to env DOCUMENT_MAX_CHUNK_SIZE)
            overlap_size: Number of characters to overlap between chunks
        """
        self.max_chunk_size = max_chunk_size or int(os.getenv('DOCUMENT_MAX_CHUNK_SIZE', '10000'))
        self.overlap_size = overlap_size
        log.info(f"DocumentProcessor initialized: max_chunk_size={self.max_chunk_size}")

    async def process_document(
        self,
        file_path: str,
        document_type: str
    ) -> Dict[str, any]:
        """
        Process a document and extract text.

        Args:
            file_path: Path to the document file
            document_type: Type of document ('pdf', 'html', 'docx')

        Returns:
            Dict with keys:
                - text: Full extracted text
                - chunks: List of text chunks
                - content_hash: SHA-256 hash of content
                - summary: First 500 chars
                - word_count: Number of words
                - page_count: Number of pages (PDF only)

        Raises:
            ProcessingError: If processing fails
        """
        try:
            file_path_obj = Path(file_path)
            if not file_path_obj.exists():
                raise ProcessingError(f"File not found: {file_path}")

            log.info(f"Processing document: {file_path} (type: {document_type})")

            # Extract text based on document type
            if document_type.lower() == 'pdf':
                text, page_count = await self._extract_pdf(file_path)
            elif document_type.lower() == 'html':
                text = await self._extract_html(file_path)
                page_count = None
            elif document_type.lower() == 'docx':
                text = await self._extract_docx(file_path)
                page_count = None
            else:
                raise ProcessingError(f"Unsupported document type: {document_type}")

            # Clean text
            text = self._clean_text(text)

            # Generate content hash
            content_hash = self._generate_hash(text)

            # Chunk text
            chunks = self._chunk_text(text, self.max_chunk_size, self.overlap_size)

            # Generate summary
            summary = self._generate_summary(text, max_length=500)

            # Count words
            word_count = len(text.split())

            result = {
                'text': text,
                'chunks': chunks,
                'content_hash': content_hash,
                'summary': summary,
                'word_count': word_count,
                'page_count': page_count,
                'chunk_count': len(chunks)
            }

            log.info(
                f"Document processed: {word_count} words, {len(chunks)} chunks, "
                f"hash={content_hash[:16]}..."
            )

            return result

        except ProcessingError:
            raise
        except Exception as e:
            log.error(f"Failed to process document {file_path}: {e}", exc_info=True)
            raise ProcessingError(f"Document processing failed: {e}")

    async def _extract_pdf(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text from PDF document.

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (extracted_text, page_count)

        Raises:
            ProcessingError: If extraction fails
        """
        try:
            text_parts = []
            page_count = 0

            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                log.debug(f"PDF has {page_count} pages")

                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    except Exception as e:
                        log.warning(f"Failed to extract page {page_num}: {e}")
                        # Continue with other pages

            if not text_parts:
                raise ProcessingError("No text could be extracted from PDF")

            text = '\n\n'.join(text_parts)
            log.debug(f"Extracted {len(text)} characters from PDF")

            return text, page_count

        except ProcessingError:
            raise
        except Exception as e:
            raise ProcessingError(f"PDF extraction failed: {e}")

    async def _extract_html(self, file_path: str) -> str:
        """
        Extract text from HTML document.

        Args:
            file_path: Path to HTML file

        Returns:
            Extracted text

        Raises:
            ProcessingError: If extraction fails
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style", "nav", "header", "footer"]):
                script.decompose()

            # Get text
            text = soup.get_text(separator='\n')

            if not text.strip():
                raise ProcessingError("No text could be extracted from HTML")

            log.debug(f"Extracted {len(text)} characters from HTML")

            return text

        except ProcessingError:
            raise
        except Exception as e:
            raise ProcessingError(f"HTML extraction failed: {e}")

    async def _extract_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX document.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text

        Raises:
            ProcessingError: If extraction fails
        """
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            if not text_parts:
                raise ProcessingError("No text could be extracted from DOCX")

            text = '\n\n'.join(text_parts)
            log.debug(f"Extracted {len(text)} characters from DOCX")

            return text

        except ProcessingError:
            raise
        except Exception as e:
            raise ProcessingError(f"DOCX extraction failed: {e}")

    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text.

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        # Normalize whitespace
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)

        # Collapse multiple newlines
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')

        # Collapse multiple spaces
        while '  ' in text:
            text = text.replace('  ', ' ')

        return text.strip()

    def _chunk_text(
        self,
        text: str,
        max_size: int,
        overlap: int = 200
    ) -> List[str]:
        """
        Split text into overlapping chunks.

        Args:
            text: Text to chunk
            max_size: Maximum characters per chunk
            overlap: Number of characters to overlap

        Returns:
            List of text chunks
        """
        if len(text) <= max_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            # Find end position
            end = start + max_size

            # If not the last chunk, try to break at sentence or paragraph
            if end < len(text):
                # Look for paragraph break
                para_break = text.rfind('\n\n', start, end)
                if para_break > start + max_size // 2:
                    end = para_break + 2

                # Otherwise look for sentence break
                else:
                    sentence_break = max(
                        text.rfind('. ', start, end),
                        text.rfind('! ', start, end),
                        text.rfind('? ', start, end)
                    )
                    if sentence_break > start + max_size // 2:
                        end = sentence_break + 2

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # Move start position with overlap
            start = end - overlap if end < len(text) else end

        log.debug(f"Split text into {len(chunks)} chunks")
        return chunks

    def _generate_hash(self, text: str) -> str:
        """
        Generate SHA-256 hash of text for deduplication.

        Args:
            text: Text to hash

        Returns:
            Hex string of hash
        """
        return hashlib.sha256(text.encode('utf-8')).hexdigest()

    def _generate_summary(self, text: str, max_length: int = 500) -> str:
        """
        Generate a summary from text (currently just truncates).

        Args:
            text: Text to summarize
            max_length: Maximum summary length

        Returns:
            Summary text
        """
        if len(text) <= max_length:
            return text

        # Truncate at sentence boundary if possible
        truncated = text[:max_length]
        last_period = truncated.rfind('. ')

        if last_period > max_length // 2:
            return truncated[:last_period + 1]
        else:
            return truncated + '...'

    def validate_file(self, file_path: str, document_type: str) -> bool:
        """
        Validate that a file exists and matches expected type.

        Args:
            file_path: Path to file
            document_type: Expected document type

        Returns:
            bool: True if valid

        Raises:
            ProcessingError: If validation fails
        """
        file_path_obj = Path(file_path)

        if not file_path_obj.exists():
            raise ProcessingError(f"File not found: {file_path}")

        if not file_path_obj.is_file():
            raise ProcessingError(f"Not a file: {file_path}")

        # Check file extension matches type
        extension = file_path_obj.suffix.lower().lstrip('.')
        if extension != document_type.lower():
            log.warning(
                f"File extension '{extension}' does not match "
                f"document type '{document_type}'"
            )

        return True


# Global processor instance
_document_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """
    Get or create the global DocumentProcessor instance.

    Returns:
        DocumentProcessor instance
    """
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
